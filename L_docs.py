from docx import Document
 
doc = Document('whistle blowing.docx')
print(doc)
for para in doc.paragraphs:
	print(para.text)
# import docx

# def read_docx(file_path):
#     doc = docx.Document(file_path)
#     full_text = []
    
#     for paragraph in doc.paragraphs:
#         full_text.append(paragraph.text)
        
#     return '\n'.join(full_text)

# input_docx = 'whistle blowing.docx'
# text = read_docx(input_docx)
# print(f'Extracted text from {input_docx}:\n{text}')